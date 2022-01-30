from tkinter import Label, messagebox, Scale, StringVar, filedialog
import tkinter as tk
from docx import Document
import csv


def choose_file():
    doc_name = tk.filedialog.askopenfilename()
    text_line_var.set(doc_name) 


#opens docx document and passes it to prep
def startup():
    try:
        doc_name = text_line.get()
        document = Document(doc_name)
    except: 
        tk.messagebox.showerror(title="Wrong Format", message="Please choose .docx")
    
    prepData(document)


def choose_csv_list():
    csv_name = tk.filedialog.askopenfilename()
    csv_line_var.set(csv_name)

    #get dict keys from csv for gui logic to have user choose which keys to include
    #todo implement gui options 
    #problem: needs a dynamic amount of checkboxes, or entry widgets for user to add keys from csv
    # i'll keed the initial thoughts in but won't follow that train of though for now as it's not necessary in the intended use case 
    # with open(csv_name, encoding='utf-8') as csvfile:
    #     names_csv_obj = csv.DictReader(csvfile)
    
    # names_indices = names_csv_obj.fieldnames
    # if names_indices:
    #     #create save btn
    #     for idx in names_indices:
    #         chk = tk.CheckButton(root, text="Speichern", command=lambda:savefile(qr_img), height=2, width=15)
    #     save_btn.grid(columnspan=2, column=2, row=5, pady=5)

# prep CSV Data and call clean, receives document from startup()
def prepData(document):
   
    result = [p.text for p in document.paragraphs]

    #try block not a real type check; only tells user what to do if run fails
    try:
    #get DictReaderObj
        csv_name = csv_line_var.get()
        with open(csv_name, encoding='utf-8') as csvfile:
            names_csv_obj = csv.DictReader(csvfile)

            #make dict out of DictReaderObj
            names_list = []
            for x in names_csv_obj:
                names_list.append(x)

            #make list out of dict
            #todo: eventually change to dynamic function that appends however many keys to list in extra function
            #needs refactoring when its not just names anymore 
            clean_names_list = []
            for row in names_list:
                clean_names_list.append(row['Vorname'])
                clean_names_list.append(row['Nachname'])
        
        #passes the current run for recursion exit in clean()
        clean(result,clean_names_list, 0)
    except:
        tk.messagebox.showerror(title="Wrong Format", message="Please choose .csv")

#clean document function, compare nameslist to text, do it twice to catch everything
def clean(text, clean_names_list, run):
    cleaned = Document()
    if run <2:
        for par in text:
            if any(name in par for name in clean_names_list):
                for name in clean_names_list:
                    if par.find(name) != -1:
                        new_sentence = par.replace(name, name[0]+'.')
                        cleaned.add_paragraph(new_sentence)
            else:
                cleaned.add_paragraph(par)
        
        cleaned_result = [p.text for p in cleaned.paragraphs]
        run += 1
        clean(cleaned_result, clean_names_list, run)
    else:
        final_result = text
        end(final_result)


#clean duplicate entries from list due to first- lastname idx issue, create doc and save doc
def end(final):
    finished_document = Document()
    for idx, f in enumerate(final):
        try:
            if final[idx] == final[idx+1]:
                pass
            else:
                finished_document.add_paragraph(f)
        except:
            pass
    try:
        filename = tk.filedialog.asksaveasfilename(defaultextension=".docx")
        finished_document.save(filename)

        tk.messagebox.showinfo(title="successfully saved", message=f"saved as {filename}")
        print(f"succesfully saved {filename}")
    except: 
        print("you broke something")   


#start tkinter 
root = tk.Tk()
root.grid_rowconfigure(0, weight=1)
root.grid_columnconfigure(0, weight=1)
canvas = tk.Canvas(root)
canvas.grid(columnspan=6, rowspan=7)

#outer layers
top_layer= tk.Label(root)
top_layer.grid(columnspan=6, column=0, row=0)
btm_layer = tk.Label(root)
btm_layer.grid(columnspan=6, column=0, row=7)
left_layer = tk.Label(root)
left_layer.grid(rowspan=7, column=0, row=0)
right_layer = tk.Label(root)
right_layer.grid(rowspan=7, column=5, row=0)

#Title
title = tk.Label(root, text="Anon-Docs")
title.grid(columnspan=4, column=1, row=0)

#instructions
instructions = tk.Label(root, text="""Clean names from a docx document, using a csv-list
!! CSV List needs first row with keys Vorname and Nachname as those keys are hardcoded
Intended use with Konfetti-generated participants list""")
instructions.grid(columnspan=4, column=1, row=1)

#CSV path box
csv_line_var = tk.StringVar()
csv_line = tk.Entry(root, width=45, textvariable=csv_line_var)
csv_line.grid(columnspan=2, column=2, row=2, sticky="w", padx=20, pady=20)
#CSV path box instruction
csv_text= tk.StringVar()
csv_text.set("Datensatz:")
csv_info=Label(root, textvariable=csv_text)
csv_info.grid(column=1, row=2)

#choose csv button
csv_btn = tk.Button(root, command=lambda:choose_csv_list(), text="Choose CSV", height=2, width=15, pady=5)
csv_btn.grid(columnspan=2, column=4, row=2)

#document path box
text_line_var = tk.StringVar()
text_line = tk.Entry(root, width=45, textvariable=text_line_var)
text_line.grid(columnspan=2, column=2, row=4, sticky="w", padx=20, pady=20)
#document path box instruction
label_text= tk.StringVar()
label_text.set("Dokument:")
label_info=Label(root, textvariable=label_text)
label_info.grid(column=1, row=4)

#choose docx button
docx_btn = tk.Button(root, command=lambda:choose_file(), text="Choose Docx", height=2, width=15, pady=5)
docx_btn.grid(columnspan=2, column=4, row=4)

#clean doc button
clean_btn = tk.Button(root, command=lambda:startup(), text="Clean and Save", height=2, width=15, pady=5)
clean_btn.grid(columnspan=2, column=2, row=6)


root.mainloop()
